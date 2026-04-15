"""
Resume → Word (.docx) export servisi.
python-docx library ishlatiladi (pip install python-docx).
"""
from io import BytesIO
from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


# ─── Rang sozlamalari (shablon bo'yicha) ─────────────────────────────────────
TEMPLATE_COLORS = {
    'minimal':      {'accent': '2E2E2E', 'light': 'F5F5F5'},
    'professional': {'accent': '1B4F8A', 'light': 'EBF2FB'},
    'modern':       {'accent': '0F7B6C', 'light': 'E8F5F3'},
    'creative':     {'accent': '7B2FBE', 'light': 'F3E8FF'},
    'academic':     {'accent': '8B1A1A', 'light': 'FBE8E8'},
}


def _hex_to_rgb(hex_color: str):
    h = hex_color.lstrip('#')
    return tuple(int(h[i:i+2], 16) for i in (0, 2, 4))


def _set_heading_color(run, hex_color: str):
    r, g, b = _hex_to_rgb(hex_color)
    run.font.color.rgb = RGBColor(r, g, b)


def _add_horizontal_rule(doc, hex_color: str = '000000'):
    """Gorizontal chiziq qo'shadi."""
    p = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '6')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), hex_color)
    pBdr.append(bottom)
    pPr.append(pBdr)
    p.paragraph_format.space_after = Pt(4)
    return p


def _section_heading(doc, text: str, accent_color: str):
    """Bo'lim sarlavhasi — katta, rangli, chiziq bilan."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run(text.upper())
    run.bold = True
    run.font.size = Pt(10)
    _set_heading_color(run, accent_color)
    _add_horizontal_rule(doc, accent_color)


def _entry_header(doc, left: str, right: str, bold_left: bool = True):
    """Ish/ta'lim sarlavhasi — chap va o'ngda."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(0)

    run_left = p.add_run(left)
    run_left.bold = bold_left
    run_left.font.size = Pt(11)

    if right:
        tab = p.add_run('\t')
        run_right = p.add_run(right)
        run_right.font.size = Pt(10)
        run_right.font.color.rgb = RGBColor(100, 100, 100)
        # Tab stop o'ngga
        pPr = p._p.get_or_add_pPr()
        tabs = OxmlElement('w:tabs')
        tab_el = OxmlElement('w:tab')
        tab_el.set(qn('w:val'), 'right')
        tab_el.set(qn('w:pos'), '9360')  # US Letter content width
        tabs.append(tab_el)
        pPr.append(tabs)
    return p


def generate_word(resume) -> bytes:
    """
    Resume obyektidan .docx fayl yaratadi.
    Returns: bytes (docx file content)
    """
    data = resume.data
    personal = data.get('personal', {})
    template_id = resume.template_id or 'minimal'
    colors = TEMPLATE_COLORS.get(template_id, TEMPLATE_COLORS['minimal'])
    accent = colors['accent']

    doc = Document()

    # ─── Sahifa sozlamalari ──────────────────────────────────────────────────
    section = doc.sections[0]
    section.page_width = Cm(21)   # A4
    section.page_height = Cm(29.7)
    section.left_margin = Cm(2)
    section.right_margin = Cm(2)
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)

    # ─── Standart uslub ─────────────────────────────────────────────────────
    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(10.5)

    # ═══════════════════════════════════════════════════════════════════════════
    # HEADER — Ism va kontakt
    # ═══════════════════════════════════════════════════════════════════════════
    name_p = doc.add_paragraph()
    name_p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    name_run = name_p.add_run(personal.get('fullName', ''))
    name_run.bold = True
    name_run.font.size = Pt(22)
    _set_heading_color(name_run, accent)
    name_p.paragraph_format.space_after = Pt(2)

    # Kontakt qatori
    contact_parts = []
    if personal.get('email'):
        contact_parts.append(personal['email'])
    if personal.get('phone'):
        contact_parts.append(personal['phone'])
    if personal.get('city'):
        contact_parts.append(personal['city'])
    if personal.get('linkedin'):
        contact_parts.append(personal['linkedin'])
    if personal.get('github'):
        contact_parts.append(personal['github'])
    if personal.get('portfolio'):
        contact_parts.append(personal['portfolio'])

    if contact_parts:
        contact_p = doc.add_paragraph(' · '.join(contact_parts))
        contact_p.runs[0].font.size = Pt(9.5)
        contact_p.runs[0].font.color.rgb = RGBColor(90, 90, 90)
        contact_p.paragraph_format.space_after = Pt(2)

    _add_horizontal_rule(doc, accent)

    # ─── Summary ────────────────────────────────────────────────────────────
    summary_text = ''
    if resume.ai_content and resume.ai_content.get('summary'):
        summary_text = resume.ai_content['summary']
    elif personal.get('summary'):
        summary_text = personal['summary']

    if summary_text:
        _section_heading(doc, 'Qisqacha tavsif', accent)
        p = doc.add_paragraph(summary_text)
        p.paragraph_format.space_after = Pt(4)

    # ─── Ish tajribasi ──────────────────────────────────────────────────────
    experience = data.get('experience', [])
    if experience:
        _section_heading(doc, 'Ish tajribasi', accent)
        ai_exp = {}
        if resume.ai_content and resume.ai_content.get('experience'):
            for item in resume.ai_content['experience']:
                ai_exp[item['index']] = item['description']

        for i, exp in enumerate(experience):
            end_date = 'Hozirgi' if exp.get('isCurrent') else exp.get('endDate', '')
            date_range = f"{exp.get('startDate', '')} – {end_date}"

            _entry_header(doc, exp.get('position', ''), date_range)
            sub_p = doc.add_paragraph(exp.get('company', ''))
            sub_p.runs[0].italic = True
            sub_p.runs[0].font.size = Pt(10)
            sub_p.runs[0].font.color.rgb = RGBColor(90, 90, 90)
            sub_p.paragraph_format.space_after = Pt(2)

            desc_text = ai_exp.get(i, exp.get('description', ''))
            if desc_text:
                desc_p = doc.add_paragraph(desc_text)
                desc_p.paragraph_format.left_indent = Cm(0.3)
                desc_p.paragraph_format.space_after = Pt(6)

    # ─── Ta'lim ─────────────────────────────────────────────────────────────
    education = data.get('education', [])
    if education:
        _section_heading(doc, 'Ta\'lim', accent)
        for edu in education:
            end_year = edu.get('endYear', '')
            date_range = f"{edu.get('startYear', '')} – {end_year or 'Hozir'}"
            degree_field = f"{edu.get('degree', '')}, {edu.get('field', '')}"
            _entry_header(doc, edu.get('institution', ''), date_range)
            sub_p = doc.add_paragraph(degree_field)
            sub_p.runs[0].italic = True
            sub_p.runs[0].font.size = Pt(10)
            sub_p.runs[0].font.color.rgb = RGBColor(90, 90, 90)
            sub_p.paragraph_format.space_after = Pt(6)

    # ─── Ko'nikmalar ────────────────────────────────────────────────────────
    skills = data.get('skills', [])
    if skills:
        _section_heading(doc, 'Ko\'nikmalar', accent)
        skills_p = doc.add_paragraph(' · '.join(skills))
        skills_p.paragraph_format.space_after = Pt(4)

    # ─── Tillar ─────────────────────────────────────────────────────────────
    languages = data.get('languages', [])
    if languages:
        _section_heading(doc, 'Tillar', accent)
        lang_items = [f"{l.get('name', '')} ({l.get('level', '')})" for l in languages]
        lang_p = doc.add_paragraph(' · '.join(lang_items))
        lang_p.paragraph_format.space_after = Pt(4)

    # ─── Sertifikatlar ──────────────────────────────────────────────────────
    certificates = data.get('certificates', [])
    if certificates:
        _section_heading(doc, 'Sertifikatlar', accent)
        for cert in certificates:
            p = doc.add_paragraph()
            run = p.add_run(cert.get('name', ''))
            run.bold = True
            run.font.size = Pt(10.5)
            p.add_run(f" — {cert.get('issuer', '')} ({cert.get('date', '')})")
            p.paragraph_format.space_after = Pt(3)

    # ─── Bytes sifatida qaytarish ────────────────────────────────────────────
    buf = BytesIO()
    doc.save(buf)
    return buf.getvalue()
